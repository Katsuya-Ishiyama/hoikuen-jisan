from datetime import date
import re
from typing import Optional, Tuple
from docx import Document
from dataclasses import dataclass
from pathlib import Path

from utils import convert_reiwa_to_year


def extract_target_month(line: str) -> Tuple[int, int]:
    year_r_str, month_str = re.findall(r"令和.*(\d+).*年度\s*(\d+).*月.*食事持参予定表", line)[0]
    year = convert_reiwa_to_year(int(year_r_str))
    month = int(month_str)
    return year, month


def extract_entry_date(line: str) -> date:
    year_r_str, month_str, day_str = re.findall(r"記入日　令和\s*(\d+)\s*年\s*(\d+)\s*月\s*(\d+)\s*", line)[0]
    year = convert_reiwa_to_year(int(year_r_str))
    month = int(month_str)
    day = int(day_str)
    return date(year, month, day)


@dataclass
class Child:
    class_name: str
    child_name: str


@dataclass
class Jisan:
    date: date
    lunch: Optional[Tuple[str, ...]]
    snack: Optional[Tuple[str, ...]]


def extract_child(line: str) -> Child:
    _class, _child = re.findall(r"\s*(\S+)\s*クラス　園児氏名\s*(.*)", line)[0]
    return Child(class_name=_class.strip(), child_name=_child.strip())


class DayNotFoundError(Exception):
    pass


class JisanData(object):
    target_month_paragraph_index = 0
    entry_date_paragraph_index = 1
    child_paragraph_index = 2

    day_cell_index = 0
    lunch_cell_index = 4
    snack_cell_index = 5

    table_body_row_index = 2

    def __init__(self, path: str):
        self.path = Path(path)
        self.doc = Document(self.path)
        self.tables = self.doc.tables

        _paragraphs = self.doc.paragraphs
        self._target_month = extract_target_month(_paragraphs[self.target_month_paragraph_index].text)
        self.entry_date = extract_entry_date(_paragraphs[self.entry_date_paragraph_index].text)
        self.child = extract_child(_paragraphs[self.child_paragraph_index].text)

        self.day_pattern = re.compile(r"\s*(\d+).*")

    @property
    def target_month(self):
        _year, _month = self._target_month
        if 1 <= _month <= 3:
            return (_year + 1), _month
        else:
            return _year, _month

    def extract_jisan_from_tables(self) -> Tuple[Jisan, ...]:
        jisan_list = []
        _year, _month = self.target_month
        for table in self.doc.tables:
            for row in table.rows[self.table_body_row_index:]:
                cells = row.cells

                try:
                    day = self._extract_jisan_day(cells[self.day_cell_index].text)
                except DayNotFoundError:
                    continue

                jisan = Jisan(
                    date=date(_year, _month, day),
                    lunch=self._extract_jisan_item(cells[self.lunch_cell_index].text),
                    snack=self._extract_jisan_item(cells[self.snack_cell_index].text)
                )
                jisan_list.append(jisan)
        return tuple(jisan_list)

    def _extract_jisan_day(self, day_str: str):
        _day_str = day_str.strip()
        if not self.day_pattern.match(_day_str):
            raise DayNotFoundError
        return int(self.day_pattern.findall(_day_str)[0])

    def _extract_jisan_item(self, jisan_str: str):
        _jisan = jisan_str.strip("\n")
        if _jisan == "":
            return None
        elif isinstance(_jisan, str):
            return (_jisan,)
        else:
            return tuple(_jisan)
